from io import BytesIO

import requests

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt
from docx.oxml.ns import qn


class DocxExporter:

    def __init__(self, blog):

        self.blog = blog

        self.document = Document()

        self._configure_document()

    # --------------------------------------------------
    # Document Configuration
    # --------------------------------------------------

    def _configure_document(self):

        section = self.document.sections[0]

        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)

        style = self.document.styles["Normal"]

        style.font.name = "Calibri"

        style._element.rPr.rFonts.set(
            qn("w:eastAsia"),
            "Calibri",
        )

        style.font.size = Pt(11)

    # --------------------------------------------------
    # Export Entry
    # --------------------------------------------------

    def export(self):

        self.add_title()

        self.add_description()

        self.add_metadata()

        self.add_hero()

        for chapter in self.blog.chapters.all():

            self.add_chapter(chapter)

        self.add_key_takeaways()

        self.add_faq()

        self.add_resources()

        self.add_conclusion()

        self.add_call_to_action()

        self.add_footer()

        return self.document

    # --------------------------------------------------
    # Helpers
    # --------------------------------------------------

    def heading(self, text, level=1):

        self.document.add_heading(text, level=level)

    def paragraph(self, text="", bold=False, italic=False):

        p = self.document.add_paragraph()

        run = p.add_run(str(text))

        run.bold = bold

        run.italic = italic

        return p

    def bullet(self, text):

        self.document.add_paragraph(
            str(text),
            style="List Bullet",
        )

    def numbered(self, text):

        self.document.add_paragraph(
            str(text),
            style="List Number",
        )

    # --------------------------------------------------
    # Image Helper
    # --------------------------------------------------

    def add_image(self, image_json, width=6):

        if not image_json:

            return

        url = (
            image_json.get("path")
            or image_json.get("url")
            or image_json.get("image_url")
        )

        if not url:

            return

        try:

            response = requests.get(
                url,
                timeout=15,
            )

            response.raise_for_status()

            image = BytesIO(
                response.content
            )

            self.document.add_picture(
                image,
                width=Inches(width),
            )

            caption = image_json.get(
                "alt_text"
            )

            if caption:

                p = self.document.add_paragraph()

                p.alignment = (
                    WD_PARAGRAPH_ALIGNMENT.CENTER
                )

                run = p.add_run(caption)

                run.italic = True

                run.font.size = Pt(9)

        except Exception:

            pass

    # --------------------------------------------------
    # Cover Page
    # --------------------------------------------------

    def add_title(self):

        heading = self.document.add_heading(
            self.blog.title,
            level=1,
        )

        heading.alignment = (
            WD_PARAGRAPH_ALIGNMENT.CENTER
        )

    def add_description(self):

        p = self.document.add_paragraph()

        p.alignment = (
            WD_PARAGRAPH_ALIGNMENT.CENTER
        )

        run = p.add_run(
            self.blog.description
        )

        run.italic = True

    def add_metadata(self):

        table = self.document.add_table(
            rows=4,
            cols=2,
        )

        table.style = "Table Grid"

        table.cell(0, 0).text = "Category"
        table.cell(0, 1).text = self.blog.category

        table.cell(1, 0).text = "Language"
        table.cell(1, 1).text = self.blog.language

        table.cell(2, 0).text = "Tone"
        table.cell(2, 1).text = self.blog.tone

        table.cell(3, 0).text = "Created"

        table.cell(3, 1).text = (
            self.blog.created_at.strftime(
                "%d %B %Y"
            )
        )

        self.document.add_paragraph()

    def add_hero(self):

        hero = self.blog.hero or {}

        self.add_image(
            hero.get("image"),
        )

        subtitle = hero.get(
            "subtitle"
        )

        if subtitle:

            self.paragraph(
                subtitle,
                italic=True,
            )

        summary = hero.get(
            "summary"
        )

        if summary:

            self.paragraph(summary)
    # --------------------------------------------------
    # Chapter
    # --------------------------------------------------

    def add_chapter(self, chapter):

        self.heading(
            chapter.title,
            level=2,
        )

        self.add_image(
            chapter.image
        )

        self.add_paragraphs(chapter)

        self.add_bullet_points(chapter)

        self.add_numbered_steps(chapter)

        self.add_tips(chapter)

        self.add_warnings(chapter)

        self.add_notes(chapter)

        self.add_examples(chapter)

        self.add_quotes(chapter)

        self.add_code(chapter)

        self.add_tables(chapter)

        self.document.add_page_break()


    # --------------------------------------------------
    # Paragraphs
    # --------------------------------------------------

    def add_paragraphs(self, chapter):

        for paragraph in chapter.paragraphs:

            self.paragraph(paragraph)


    # --------------------------------------------------
    # Bullet Points
    # --------------------------------------------------

    def add_bullet_points(self, chapter):

        if not chapter.bullet_points:

            return

        self.heading(
            "Key Points",
            level=3,
        )

        for point in chapter.bullet_points:

            self.bullet(point)


    # --------------------------------------------------
    # Numbered Steps
    # --------------------------------------------------

    def add_numbered_steps(self, chapter):

        if not chapter.numbered_steps:

            return

        self.heading(
            "Steps",
            level=3,
        )

        for step in chapter.numbered_steps:

            if isinstance(step, dict):

                title = step.get(
                    "title",
                    "",
                )

                description = step.get(
                    "description",
                    "",
                )

                self.numbered(title)

                if description:

                    self.paragraph(description)

            else:

                self.numbered(step)
    # --------------------------------------------------
    # Key Takeaways
    # --------------------------------------------------

    def add_key_takeaways(self):

        if not self.blog.key_takeaways:
            return

        self.document.add_page_break()

        self.heading(
            "Key Takeaways",
            level=2,
        )

        for takeaway in self.blog.key_takeaways:

            self.bullet(takeaway)
    # --------------------------------------------------
    # Tips
    # --------------------------------------------------

    def add_tips(self, chapter):

        if not chapter.tips:

            return

        self.heading(
            "Tips",
            level=3,
        )

        for tip in chapter.tips:

            p = self.document.add_paragraph()

            run = p.add_run("💡 ")

            run.bold = True

            p.add_run(tip)
    # --------------------------------------------------
    # Warnings
    # --------------------------------------------------

    def add_warnings(self, chapter):

        if not chapter.warnings:

            return

        self.heading(
            "Things to Watch Out For",
            level=3,
        )

        for warning in chapter.warnings:

            p = self.document.add_paragraph()

            run = p.add_run("⚠ ")

            run.bold = True

            p.add_run(warning)
    # --------------------------------------------------
    # Notes
    # --------------------------------------------------

    def add_notes(self, chapter):

        if not chapter.notes:

            return

        self.heading(
            "Notes",
            level=3,
        )

        for note in chapter.notes:

            self.bullet(note)
    # --------------------------------------------------
    # Examples
    # --------------------------------------------------

    def add_examples(self, chapter):

        if not chapter.examples:
            return

        self.heading(
            "Examples",
            level=3,
        )

        for example in chapter.examples:

            if isinstance(example, dict):

                title = example.get(
                    "title",
                    "",
                )

                description = example.get(
                    "description",
                    "",
                )

                if title:

                    self.paragraph(
                        title,
                        bold=True,
                    )

                if description:

                    self.paragraph(
                        description,
                    )

            else:

                self.paragraph(
                    str(example),
                )
        # --------------------------------------------------
    # Quotes
    # --------------------------------------------------

    def add_quotes(self, chapter):

        if not chapter.quotes:

            return

        self.heading(
            "Quotes",
            level=3,
        )

        for quote in chapter.quotes:

            p = self.document.add_paragraph()

            p.style = "Quote"

            p.add_run(quote)
    # --------------------------------------------------
    # Code
    # --------------------------------------------------

    def add_code(self, chapter):

        if not chapter.code_examples:

            return

        self.heading(
            "Code Examples",
            level=3,
        )

        for code in chapter.code_examples:

            language = code.get(
                "language",
                "",
            )

            snippet = code.get(
                "code",
                "",
            )

            if language:

                self.paragraph(
                    language,
                    bold=True,
                )

            p = self.document.add_paragraph()

            run = p.add_run(snippet)

            run.font.name = "Consolas"

            run.font.size = Pt(10)
    # --------------------------------------------------
    # Tables
    # --------------------------------------------------

    def add_tables(self, chapter):

        if not chapter.tables:

            return

        self.heading(
            "Tables",
            level=3,
        )

        for table_data in chapter.tables:

            headers = table_data.get(
                "headers",
                [],
            )

            rows = table_data.get(
                "rows",
                [],
            )

            if not headers:

                continue

            table = self.document.add_table(

                rows=1,

                cols=len(headers),

            )

            table.style = "Table Grid"

            hdr = table.rows[0].cells

            for i, header in enumerate(headers):

                hdr[i].text = str(header)

            for row in rows:

                cells = table.add_row().cells

                for i, value in enumerate(row):

                    if i < len(cells):

                        cells[i].text = str(value)

            self.document.add_paragraph()
    # --------------------------------------------------
    # FAQ
    # --------------------------------------------------

    def add_faq(self):

        if not self.blog.faq:
            return

        self.document.add_page_break()

        self.heading(
            "Frequently Asked Questions",
            level=2,
        )

        for faq in self.blog.faq:

            question = faq.get(
                "question",
                "",
            )

            answer = faq.get(
                "answer",
                "",
            )

            self.paragraph(
                question,
                bold=True,
            )

            self.paragraph(answer)
    # --------------------------------------------------
    # Resources
    # --------------------------------------------------

    def add_resources(self):

        if not self.blog.resources:
            return

        self.heading(
            "Resources",
            level=2,
        )

        for resource in self.blog.resources:

            if isinstance(resource, dict):

                title = resource.get(
                    "title",
                    "",
                )

                url = resource.get(
                    "url",
                    "",
                )

                if title:

                    self.paragraph(
                        title,
                        bold=True,
                    )

                if url:

                    self.paragraph(url)

            else:

                self.paragraph(resource)
    # --------------------------------------------------
    # Conclusion
    # --------------------------------------------------

    def add_conclusion(self):

        if not self.blog.conclusion:
            return

        self.document.add_page_break()

        self.heading(
            "Conclusion",
            level=2,
        )

        self.add_image(
            self.blog.conclusion.get(
                "image"
            )
        )

        paragraphs = self.blog.conclusion.get(
            "paragraphs",
            [],
        )

        for paragraph in paragraphs:

            self.paragraph(paragraph)
    # --------------------------------------------------
    # Call To Action
    # --------------------------------------------------

    def add_call_to_action(self):

        if not self.blog.call_to_action:
            return

        self.heading(
            "Call To Action",
            level=2,
        )

        if isinstance(
            self.blog.call_to_action,
            dict,
        ):

            title = self.blog.call_to_action.get(
                "title",
                "",
            )

            description = self.blog.call_to_action.get(
                "description",
                "",
            )

            button = self.blog.call_to_action.get(
                "button_text",
                "",
            )

            if title:

                self.paragraph(
                    title,
                    bold=True,
                )

            if description:

                self.paragraph(description)

            if button:

                self.paragraph(
                    f"Action: {button}",
                    italic=True,
                )

        else:

            self.paragraph(
                self.blog.call_to_action
            )
    # --------------------------------------------------
    # Footer
    # --------------------------------------------------

    def add_footer(self):

        self.document.add_page_break()

        p = self.document.add_paragraph()

        p.alignment = (
            WD_PARAGRAPH_ALIGNMENT.CENTER
        )

        run = p.add_run(
            "Generated using AI Blog Generator"
        )

        run.bold = True

        run.font.size = Pt(12)

        p = self.document.add_paragraph()

        p.alignment = (
            WD_PARAGRAPH_ALIGNMENT.CENTER
        )

        p.add_run(
            self.blog.created_at.strftime(
                "%d %B %Y"
            )
        )